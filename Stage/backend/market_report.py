from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from periodic_task.build_strategy_index import  get_stg_index_quantile
from fh_tools.fh_utils import  return_risk_analysis
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import datetime
import logging
myfont = fm.FontProperties(fname='/usr/share/fonts/simhei.ttf')
basedir = os.path.abspath(os.path.dirname(__file__))
logger = logging.getLogger()

def gen_analysis(start,end):
    stg_idx_quantile_dic, stat_df, stg_idx_mid_df = get_stg_index_quantile(start, end, do_filter=3, mgrcomp_id_2_name=True)
    table_df = return_risk_analysis(stg_idx_mid_df, freq=None).drop('截止日期')
    table_df.reset_index(inplace=True)
    table_df.rename(columns={'index':''}, inplace=True)
    text_df = stat_df.T
    text_df.reset_index(inplace=True)
    text_dict = text_df.to_dict("records")
    return {"charts_dict":stg_idx_quantile_dic,"table_df":table_df,"text_dict":text_dict}

def gen_report(start,end):
    document = Document()
    analysis_data = gen_analysis(start,end)
    table_df = analysis_data['table_df']
    for k,v in analysis_data['charts_dict'].items():
        df = v['date_idx_quantile_df']
        df.index.name = ''
        plt.rcParams["figure.figsize"] = (10, 3)
        plt.style.use('fivethirtyeight')
        ax = df.plot(fontsize=7)
        for i, l in enumerate(ax.lines):
            plt.setp(l, linewidth=1.8)
        box = ax.get_position()
        ax.set_position([box.x0, box.y0, box.width * 0.88, box.height])
        plt.legend(loc=6, prop={'size': 8, 'family': 'simhei'}, bbox_to_anchor=(1.05, 0.5))
        plt.title(k, fontproperties=myfont, size=9)
        plt.savefig("%s.png"%k,dpi=300,transparent=True)

    document.add_heading('XX季度策略表现回顾',2)
    for i in analysis_data['text_dict']:
        text = ''
        for k,v in sorted(i.items()):
            if k == 'index':
                t = '%s:'%v
                document.add_picture('%s.png' % v, width=Inches(6))
                os.remove('%s.png' %v)
            else:
                t = "%s:%s," %(k,v)
            text +=t
        document.add_paragraph(text[:-1])

    document.add_page_break()


    #增加表格
    t = document.add_table(table_df.shape[0]+1, table_df.shape[1])
    t.style = 'Table Grid'
    # add the header rows.
    for j in range(table_df.shape[-1]):
        t.cell(0,j).text = table_df.columns[j]

    # add the rest of the data frame
    for i in range(table_df.shape[0]):
        for j in range(table_df.shape[-1]):
            t.cell(i+1,j).text = str(table_df.values[i,j])
    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)


    #保存文件
    file_name = datetime.datetime.now().strftime("%Y-%m-%d_%H%M%S")
    file_path = os.path.join(basedir,file_name)
    document.save('%s.docx' %file_path)
    return "{}.docx".format(file_path),analysis_data

if __name__ == '__main__':
    x = gen_report('2017-07-1','2017-09-30')
    print(x)